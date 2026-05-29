"""Generate SP500 App — Maintainer Wiki as a Word (.docx) document."""
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
C_HEADING1  = RGBColor(0x1a, 0x4e, 0x8a)   # dark blue
C_HEADING2  = RGBColor(0x2c, 0x3e, 0x50)   # charcoal
C_HEADING3  = RGBColor(0x27, 0xae, 0x60)   # green
C_CODE_BG   = RGBColor(0xf4, 0xf4, 0xf4)   # light grey (simulated via shading)
C_WHITE     = RGBColor(0xff, 0xff, 0xff)
C_THEAD_BG  = RGBColor(0x1a, 0x4e, 0x8a)
C_ROW_ALT   = RGBColor(0xeb, 0xf0, 0xf7)
C_WARNING   = RGBColor(0xc0, 0x39, 0x2b)
C_TIP       = RGBColor(0x27, 0xae, 0x60)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Add borders to a cell. kwargs: top/bottom/left/right = True."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        if kwargs.get(side):
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "2C3E50")
            tcBorders.append(border)
    tcPr.append(tcBorders)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size  = Pt(20)
    run.font.bold  = True
    run.font.color.rgb = C_HEADING1
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A4E8A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = C_HEADING2
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.size   = Pt(12)
    run.font.bold   = True
    run.font.italic = True
    run.font.color.rgb = C_HEADING3
    return p


def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def note(text, kind="note"):
    """Callout paragraph — kind: 'note', 'warn', 'tip'."""
    colors = {"note": C_HEADING2, "warn": C_WARNING, "tip": C_TIP}
    icons  = {"note": "ℹ ", "warn": "⚠  ", "tip": "✔ "}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(icons[kind] + text)
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = colors[kind]
    return p


def code(text):
    """Monospaced code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    # Shade paragraph background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F4F4F4")
    pPr.append(shd)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def table(headers, rows, col_widths=None):
    """Styled table with coloured header row and alternating row shading."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style     = "Table Grid"

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold       = True
        cell.paragraphs[0].runs[0].font.color.rgb  = C_WHITE
        cell.paragraphs[0].runs[0].font.size        = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, C_THEAD_BG)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = C_ROW_ALT if ri % 2 == 1 else C_WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return t


def page_break():
    doc.add_page_break()


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(4)
p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("S&P 500 Stock Analysis App")
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = C_HEADING1

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Maintainer Reference Guide")
run.font.size   = Pt(16)
run.font.italic = True
run.font.color.rgb = C_HEADING2

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
run.font.size  = Pt(10)
run.font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GitHub Repository: github.com/kishor777/SP500")
run.font.size  = Pt(10)
run.font.color.rgb = C_HEADING1

doc.add_paragraph()
divider()
doc.add_paragraph()

# Contents overview
h2("Document Contents")
table(
    ["Section", "What it covers"],
    [
        ["1 · Architecture & Features",  "How the app works, pages, data flow, database schema"],
        ["2 · Configuration Reference",  "Every constant in config.py explained in plain English"],
        ["3 · Deployment Guide",         "Install on a new Windows PC, restore DB, run the app"],
    ],
    col_widths=[7, 9.5],
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — ARCHITECTURE & FEATURES
# ══════════════════════════════════════════════════════════════════════════════

h1("1 · Architecture & Features")

h2("1.1  System Overview")
body(
    "The app is a Flask web application running on port 5050. It reads stock data from a "
    "local MySQL database, fetches live prices from Yahoo Finance every 5 minutes during "
    "market hours, and scores all 503 S&P 500 stocks using a six-signal recommendation "
    "engine. A separate sentiment module analyses social media and news via FinBERT NLP."
)

h3("High-Level Architecture")
code(
    "Browser (port 5050)\n"
    "    │\n"
    "    ▼\n"
    "Flask App  ─────────────────────────────────────────────────────┐\n"
    "  sp500_viewer.py                                               │\n"
    "  ├─ HTTP Routes & API endpoints                                │\n"
    "  ├─ Background Scheduler (every 5 min, market hours)           │\n"
    "  ├─ Recommendation Engine (6-signal composite scorer)          │\n"
    "  └─ Sentiment Engine  (sentiment_engine.py)                    │\n"
    "                                                                │\n"
    "  config.py ──── All constants & weights                        │\n"
    "  db_config.py ─ DB connection (credentials from .env)          │\n"
    "                                                                │\n"
    "MySQL Database ─────────────────────────────────────────────────┘\n"
    "  sp500_info · sp500_history · sp500_intraday\n"
    "  guru_holdings · guru_rules · screener_saved_filters\n"
    "  vol_trades · sentiment_scores\n"
    "                                                                   \n"
    "External APIs\n"
    "  Yahoo Finance (yfinance) ── price data, fundamentals\n"
    "  SEC EDGAR ─────────────── 13F filing XML\n"
    "  StockTwits / Yahoo RSS ─── social sentiment text\n"
    "  Reddit (PRAW) ─────────── subreddit posts\n"
    "  FinBERT (local CPU) ────── NLP sentiment scoring"
)

h2("1.2  Technology Stack")
table(
    ["Component", "Technology", "Version"],
    [
        ["Web framework",    "Flask (Python)",          "3.x"],
        ["Database",         "MySQL",                   "8.x"],
        ["Data access",      "SQLAlchemy",              "2.x"],
        ["Market data",      "yfinance",                "0.2+"],
        ["SEC filings",      "requests + XML parser",   "—"],
        ["Sentiment NLP",    "FinBERT (Hugging Face)",  "CPU inference"],
        ["Social data",      "StockTwits API / PRAW",   "—"],
        ["Charts",           "lightweight-charts (JS)", "4.x"],
        ["Tables",           "DataTables (JS)",         "1.13"],
        ["Credential mgmt",  "python-dotenv",           "1.x"],
    ],
    col_widths=[5, 6.5, 5],
)

h2("1.3  Application Pages")
table(
    ["URL", "Page Name", "Purpose"],
    [
        ["/",                "Main Table",      "All 503 S&P 500 stocks — sortable, with candlestick charts"],
        ["/screener",        "Screener",        "Filter stocks by 50+ financial metrics; save and load presets"],
        ["/gurus",           "Guru Investing",  "13F portfolio holdings for 21 top investors"],
        ["/leaderboard",     "Leaderboard",     "Ranks gurus by how well their screened picks performed"],
        ["/recommendations", "Recommendations", "AI-scored ranking of all 503 stocks with live weight sliders"],
    ],
    col_widths=[3.5, 4, 9],
)

h2("1.4  Features in Detail")

h3("Main Table  (/)")
bullet("Loads all 503 S&P 500 stocks from the sp500_info MySQL table")
bullet("Sortable columns: Name, Sector, Price, Market Cap, P/E, Dividend Yield, ROE, Revenue Growth, 1-Year Return, Analyst Rating")
bullet("Click any ticker → candlestick chart popup with Daily and 5-min bar toggle")
bullet("Chart popup includes: price chart, volume chart (colour-coded green/red), full company fundamental panel")
note("Screenshot: Main table showing dark-theme table with sector tags, price column, and chart popup open on right side.", "note")

h3("Screener  (/screener)")
body("Three tabs inside the Screener page:")
bullet("Screener tab — 50+ range sliders across 7 groups. Save named filter presets. Apply any guru's investment rules as a preset.", level=0)
bullet("Volume Analysis tab — abnormal volume detection (today > 2× 1-month average). Shows upcoming earnings dates and average volumes.", level=0)
bullet("Vol Trades tab — paper trading performance dashboard. Shows Realized P&L, Win Rate, and Today/Weekly/Monthly returns.", level=0)
body("Auto paper trading logic:")
bullet("10:00 AM ET — buys top 3 bullish stocks from Volume Analysis ($1,000 each)")
bullet("2:00 PM ET  — sells all open positions from today")
note("Screenshot: Screener page with filter sliders visible on left, matching stocks table on right, Guru preset dropdown active.", "note")

h3("Guru Investing  (/gurus)")
body("Tracks 21 top investors using quarterly 13F SEC filings:")
table(
    ["Guru", "Fund", "Style"],
    [
        ["Warren Buffett",       "Berkshire Hathaway",        "Quality Value"],
        ["Bill Ackman",          "Pershing Square",           "Concentrated Quality"],
        ["David Einhorn",        "Greenlight Capital",        "Deep Value"],
        ["Michael Burry",        "Scion Asset Management",    "Deep Value / Contrarian"],
        ["Stanley Druckenmiller","Duquesne Family Office",    "Macro Growth"],
        ["Carl Icahn",           "Icahn Capital",             "Activist Value"],
        ["Ray Dalio",            "Bridgewater Associates",    "All Weather / Risk Parity"],
        ["Dan Loeb",             "Third Point LLC",           "Event-Driven Value"],
        ["Chase Coleman",        "Tiger Global",              "High-Growth Tech"],
        ["Tom Lee",              "Fundstrat (GRNY ETF)",      "Growth & Momentum"],
        ["David Tepper",         "Appaloosa Management",      "Distressed & Macro"],
        ["Paul Tudor Jones",     "Tudor Investment Corp",     "Global Macro & Momentum"],
        ["Cathie Wood",          "ARK Investment",            "Disruptive Innovation"],
        ["Ken Griffin",          "Citadel Advisors",          "Multi-Strategy Quant"],
        ["Jim Simons",           "Renaissance Technologies",  "Quantitative Systematic"],
        ["Seth Klarman",         "Baupost Group",             "Deep Value / Margin of Safety"],
        ["Howard Marks",         "Oaktree Capital",           "Risk-Aware Value"],
        ["George Soros",         "Soros Fund Management",     "Global Macro"],
        ["Philippe Laffont",     "Coatue Management",         "Growth Technology"],
        ["Ken Fisher",           "Fisher Asset Management",   "GARP / Global Large-Cap"],
        ["Brad Gerstner",        "Altimeter Capital",         "AI & Software Growth"],
    ],
    col_widths=[5.5, 5.5, 5.5],
)
note("Screenshot: Guru page with Buffett tab selected, portfolio table showing ticker, weight%, value, and change tags (New/Added/Held/Reduced).", "note")

h3("Recommendation Engine  (/recommendations)")
body(
    "Scores all 503 stocks using a composite of 6 signals. Each signal produces a "
    "0–100 percentile rank within the S&P 500 universe. The final score is a weighted "
    "sum cached for 1 hour."
)
table(
    ["Signal", "Default Weight", "What it measures"],
    [
        ["Momentum",          "15%", "3M / 6M / 1Y price return blend — percentile ranked"],
        ["Fundamental Quality","25%", "ROE, margins, and revenue/earnings growth — percentile ranked"],
        ["Valuation",         "15%", "P/E, P/B, P/S, EV/EBITDA — inverted rank (cheaper = higher score)"],
        ["Guru Conviction",   "15%", "13F holdings: New=3pts, Added=2pts, Held=1pt, Reduced=−0.5pts"],
        ["Analyst Consensus", "15%", "Strong Buy=100, Buy=80, Hold=50, Underperform=25, Sell=5"],
        ["Market Sentiment",  "15%", "StockTwits 40% + Reddit 30% + Volume spike 20% + Momentum 10%"],
    ],
    col_widths=[4.5, 3.5, 8.5],
)
body("Score labels:")
table(
    ["Score Range", "Label"],
    [
        ["68 – 100", "Strong Buy"],
        ["57 – 67",  "Buy"],
        ["44 – 56",  "Hold"],
        ["32 – 43",  "Underperform"],
        ["0 – 31",   "Sell"],
    ],
    col_widths=[4, 12.5],
)
body("Interactive features on this page:")
bullet("Weight sliders — drag to rebalance the 6 signals; table re-ranks live in the browser (no server call)")
bullet("⚡ Run Sentiment button — triggers on-demand FinBERT sentiment scoring for Top 20 + Bottom 20 stocks")
bullet("ⓘ Info button — popup explaining every weight, threshold, and sub-signal in plain English")
note("Screenshot: Recommendations page with weight slider panel, score bars, label badges, and Sentiment column showing 'Bullish'/'Bearish' classification badges.", "note")

h2("1.5  Scheduler — Background Thread")
body(
    "A background thread runs every 5 minutes. It checks whether the NYSE market is open "
    "(weekdays 9:30–16:00 ET, excluding holidays) before taking any action."
)
table(
    ["Task", "When", "Action"],
    [
        ["Price refresh",      "Every 5 min (market hours)",  "Fetch 5-min OHLCV bars from yfinance for all 503 tickers"],
        ["History backfill",   "App startup",                 "Download 2 years of daily OHLCV if not already in DB"],
        ["Intraday backfill",  "App startup",                 "Download 60 days of 5-min bars if below threshold"],
        ["Paper buy",          "First refresh ≥ 10:00 AM ET","Buy top 3 bullish Volume Analysis stocks ($1,000 each)"],
        ["Paper sell",         "First refresh ≥ 2:00 PM ET", "Sell all open paper positions from today"],
        ["Sentiment pass",     "On demand / once daily",      "Score Top 20 + Bottom 20 recommendation tickers via FinBERT"],
    ],
    col_widths=[4, 5, 7.5],
)

h2("1.6  Database Tables")
table(
    ["Table", "Rows (approx.)", "Re-fetchable?", "Contents"],
    [
        ["sp500_info",             "503",      "Yes (yfinance)",   "Company fundamentals — price, P/E, margins, growth"],
        ["sp500_history",          "~370,000", "Yes (yfinance)",   "Daily OHLCV — 2 years per ticker"],
        ["sp500_intraday",         "varies",   "Yes (yfinance)",   "5-minute OHLCV bars — rolling 60 days"],
        ["guru_holdings",          "~5,000+",  "Yes (SEC EDGAR)",  "13F filing positions per guru per quarter"],
        ["guru_filing_meta",       "~40",      "Yes (SEC EDGAR)",  "Filing dates and fetch timestamps"],
        ["guru_rules",             "~130",     "NO — back up",     "Customised investment thresholds per guru"],
        ["screener_saved_filters", "varies",   "NO — back up",     "Named screener presets created by the user"],
        ["vol_trades",             "varies",   "NO — back up",     "Paper trading history — buy/sell records"],
        ["sentiment_scores",       "varies",   "NO — back up",     "Daily FinBERT sentiment scores per ticker"],
    ],
    col_widths=[5, 3, 3, 5.5],
)
note("Back up the four 'NO' tables regularly using: python backup.py  then commit to GitHub.", "warn")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CONFIGURATION REFERENCE
# ══════════════════════════════════════════════════════════════════════════════

h1("2 · Configuration Reference")
body(
    "All tuneable constants live in config.py. Credentials live in .env only (never committed "
    "to GitHub). This section explains every setting in plain English."
)

h2("2.1  Credentials  (.env file)")
note("Never paste real passwords into config.py or any file committed to GitHub.", "warn")
table(
    ["Variable", "Example", "What it does"],
    [
        ["DB_HOST",             "localhost",   "MySQL server address"],
        ["DB_PORT",             "3306",        "MySQL port (default 3306)"],
        ["DB_USER",             "root",        "MySQL username"],
        ["DB_PASSWORD",         "yourpassword","MySQL password — keep secret"],
        ["DB_NAME",             "sp500",       "Name of the MySQL database"],
        ["REDDIT_CLIENT_ID",    "abc123",      "From reddit.com/prefs/apps — leave blank if not used"],
        ["REDDIT_CLIENT_SECRET","xyz456",      "From reddit.com/prefs/apps — leave blank if not used"],
    ],
    col_widths=[5, 4, 7.5],
)

h2("2.2  Scheduler Settings  (SCHEDULER_CONFIG)")
table(
    ["Key", "Default", "When to change"],
    [
        ["refresh_interval_sec",   "300",  "How often price refresh runs (seconds). 300 = every 5 min."],
        ["vol_trade_buy_hour",     "10",   "ET hour for paper buy. 10 = 10:00 AM ET."],
        ["vol_trade_sell_hour",    "14",   "ET hour for paper sell. 14 = 2:00 PM ET."],
        ["intraday_batch_size",    "50",   "Tickers per yfinance API call. Reduce to 25 if timeouts occur."],
        ["intraday_backfill_days", "60",   "Days of 5-min bars to store. Increase for more chart history."],
        ["intraday_skip_threshold","1000", "Skip backfill if DB already has this many rows."],
        ["hist_batch_size",        "50",   "Tickers per batch for daily history backfill."],
        ["hist_backfill_days",     "730",  "Days of daily history (730 = 2 years)."],
    ],
    col_widths=[5, 2, 9.5],
)

h2("2.3  Cache Timeouts")
table(
    ["Constant", "Default", "Meaning"],
    [
        ["_DB_GURU_TTL",   "86400 (24h)", "Seconds before 13F guru data is re-fetched from SEC EDGAR."],
        ["_OPTIMIZED_TTL", "3600 (1h)",   "Seconds before Optimised Guru rules recalculate."],
        ["_RECO_TTL",      "3600 (1h)",   "Seconds before Recommendation Engine rescores all stocks."],
    ],
    col_widths=[4.5, 3.5, 8.5],
)

h2("2.4  Recommendation Signal Weights  (RECO_SIGNAL_WEIGHTS)")
note("These six weights must always sum to 1.0 (100%). The Recommendations page sliders let you override these live without editing this file.", "tip")
table(
    ["Signal", "Default", "What it measures"],
    [
        ["momentum",    "0.15 (15%)", "Recent price performance across 3M, 6M, and 1Y periods"],
        ["fundamental", "0.25 (25%)", "Profitability and growth quality metrics"],
        ["valuation",   "0.15 (15%)", "How cheap the stock is relative to earnings, book, and sales"],
        ["guru",        "0.15 (15%)", "How many top investors hold it and whether they are buying"],
        ["analyst",     "0.15 (15%)", "Wall Street analyst consensus rating"],
        ["sentiment",   "0.15 (15%)", "Social media and news sentiment scored by FinBERT NLP"],
    ],
    col_widths=[4, 3, 9.5],
)

h2("2.5  Score Labels  (RECO_LABEL_THRESHOLDS)")
table(
    ["Score", "Label", "When to adjust"],
    [
        ["≥ 68", "Strong Buy",   "Lower to 65 if too few Strong Buys appear"],
        ["≥ 57", "Buy",          "Lower to 52 if Buy list is too small"],
        ["≥ 44", "Hold",         "Usually leave as-is — neutral midpoint"],
        ["≥ 32", "Underperform", "Raise to 36 to widen the Underperform bucket"],
        ["< 32", "Sell",         "Lower to 25 to reduce Sell count"],
    ],
    col_widths=[2.5, 4, 10],
)

h2("2.6  Momentum Sub-weights  (RECO_MOMENTUM_WEIGHTS)")
table(
    ["Period", "Default", "Why"],
    [
        ["3-month return (m3)", "30%", "Medium-term trend confirmation"],
        ["6-month return (m6)", "40%", "Most predictive period for continuation — highest weight"],
        ["1-year return (y1)",  "30%", "Long-term trend backdrop"],
    ],
    col_widths=[5, 2.5, 9],
)
note("1-week and 1-month returns appear in the table display but are NOT used in scoring — too noisy.", "note")

h2("2.7  Fundamental Quality Fields  (RECO_FUNDAMENTAL_FIELDS)")
table(
    ["Metric (Yahoo Finance field)", "Weight", "Plain English"],
    [
        ["returnOnEquity",   "30%", "Profit generated per dollar of shareholder equity — measures management effectiveness"],
        ["grossMargins",     "20%", "Revenue minus cost of goods — measures pricing power"],
        ["operatingMargins", "20%", "Profit after operating costs — measures operational efficiency"],
        ["revenueGrowth",    "15%", "Year-over-year sales growth rate"],
        ["earningsGrowth",   "15%", "Year-over-year net income growth rate"],
    ],
    col_widths=[5.5, 2, 9],
)

h2("2.8  Valuation Fields  (RECO_VALUATION_FIELDS)")
note("Lower multiples score higher here — cheaper stocks rank better. Negative values (loss-making companies) are clamped to 0 before ranking.", "tip")
table(
    ["Metric", "Weight", "Plain English"],
    [
        ["trailingPE",                    "30%", "Share price ÷ last 12 months earnings per share"],
        ["priceToBook",                   "25%", "Share price ÷ net asset value per share"],
        ["priceToSalesTrailing12Months",  "25%", "Market cap ÷ last 12 months revenue"],
        ["enterpriseToEbitda",            "20%", "Enterprise value ÷ operating profit — standard M&A metric"],
    ],
    col_widths=[6, 2, 8.5],
)

h2("2.9  Guru Conviction Tag Points  (RECO_GURU_TAG_WEIGHT)")
table(
    ["13F Change Tag", "Points", "Meaning"],
    [
        ["new",      "+3.0", "Guru opened a brand new position — highest conviction signal"],
        ["added",    "+2.0", "Guru increased an existing position"],
        ["held",     "+1.0", "Guru maintained the position unchanged"],
        ["(unknown)","+1.0", "Tag not available — treated as held"],
        ["reduced",  "−0.5", "Guru trimmed position — mild negative signal"],
    ],
    col_widths=[4, 2, 10.5],
)

h2("2.10  Analyst Score Map  (RECO_ANALYST_SCORE_MAP)")
table(
    ["Yahoo Finance Rating", "Score"],
    [
        ["Strong Buy",   "100"],
        ["Buy",          "80"],
        ["Hold",         "50"],
        ["Underperform", "25"],
        ["Sell",         "5"],
        ["Missing",      "50 — neutral, no penalty"],
    ],
    col_widths=[6, 10.5],
)

h2("2.11  Sentiment Engine Settings  (SENTIMENT_*)")
table(
    ["Constant", "Default", "What it does"],
    [
        ["SENTIMENT_STOCKTWITS_WEIGHT", "0.40", "StockTwits sentiment within the sentiment composite"],
        ["SENTIMENT_REDDIT_WEIGHT",     "0.30", "Reddit sentiment within the sentiment composite"],
        ["SENTIMENT_VOLUME_WEIGHT",     "0.20", "Mention volume spike vs 3-day baseline"],
        ["SENTIMENT_MOMENTUM_WEIGHT",   "0.10", "Sentiment trend vs yesterday"],
        ["SENTIMENT_LOOKBACK_DAYS",     "3",    "Only analyse posts from the last N days"],
        ["SENTIMENT_MIN_POST_LEN",      "20",   "Ignore posts shorter than this (characters) — spam filter"],
        ["SENTIMENT_REFRESH_HOUR",      "12",   "ET hour after which a second daily run is allowed"],
        ["SENTIMENT_SUBREDDITS",        "5 subs","wallstreetbets, stocks, investing, SecurityAnalysis, options"],
    ],
    col_widths=[6, 2.5, 8],
)

h2("2.12  NYSE Market Calendar")
body(
    "The app knows NYSE holidays for 2025–2026 (config.py → _NYSE_HOLIDAYS). "
    "The scheduler will not fetch data on these days. To add a new year of holidays, "
    "add date entries to that list in config.py — example:"
)
code(
    "_date(2027, 1, 1),   # New Year's Day\n"
    "_date(2027, 1, 18),  # MLK Day\n"
    "_date(2027, 2, 15),  # Presidents Day\n"
    "# ... etc"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DEPLOYMENT GUIDE
# ══════════════════════════════════════════════════════════════════════════════

h1("3 · Deployment Guide")
body(
    "Complete step-by-step instructions for installing the app on a new Windows machine "
    "and restoring data from a GitHub backup."
)

h2("3.1  Prerequisites")
note("Install all of these before starting. All downloads are free.", "tip")
table(
    ["Software", "Version", "Download URL"],
    [
        ["Python",           "3.11 or newer", "https://www.python.org/downloads/"],
        ["MySQL Server",     "8.x",           "https://dev.mysql.com/downloads/mysql/"],
        ["MySQL Workbench",  "8.x (optional)","https://dev.mysql.com/downloads/workbench/"],
        ["Git",              "Latest",        "https://git-scm.com/download/win"],
    ],
    col_widths=[4, 3.5, 9],
)
note("During Python install — tick 'Add Python to PATH' before clicking Install.", "warn")
note("During MySQL install — note down the root password you set. You will need it in Step 4.", "warn")

h2("Step 1 — Get the Code from GitHub")
body("Open PowerShell and run:")
code(
    "cd C:\\\n"
    "git clone https://github.com/kishor777/SP500.git\n"
    "cd SP500"
)
body("This creates C:\\SP500\\ with all application files.")

h2("Step 2 — Install Python Packages")
code("pip install -r requirements.txt")
body("This installs Flask, SQLAlchemy, yfinance, and all other dependencies. Takes 2–5 minutes.")
note(
    "Optional — for FinBERT sentiment scoring, also run:\n"
    "pip install transformers torch\n"
    "This downloads a ~500 MB model on first use. The app works without it — stocks just get a neutral sentiment score.",
    "note"
)

h2("Step 3 — Create the Database")
body("Open MySQL Workbench and run, or paste into PowerShell:")
code(
    "mysql -u root -p -e \"\n"
    "  CREATE DATABASE sp500\n"
    "  CHARACTER SET utf8mb4\n"
    "  COLLATE utf8mb4_unicode_ci;\n"
    "\""
)

h2("Step 4 — Create Your Credentials File")
body("In C:\\SP500\\, create a file called .env (note the dot at the start):")
code("notepad C:\\SP500\\.env")
body("Paste this content, replacing values with your own:")
code(
    "DB_HOST=localhost\n"
    "DB_PORT=3306\n"
    "DB_USER=root\n"
    "DB_PASSWORD=your_mysql_password_here\n"
    "DB_NAME=sp500\n"
    "\n"
    "REDDIT_CLIENT_ID=\n"
    "REDDIT_CLIENT_SECRET="
)
note("This file must NEVER be shared or uploaded to GitHub. It is already excluded by .gitignore.", "warn")

h2("Step 5 — Create db_config.py")
body("Copy the example file and rename it:")
code("Copy-Item C:\\SP500\\db_config.example.py C:\\SP500\\db_config.py")
body("The app reads credentials from .env automatically — no need to edit db_config.py.")

h2("Step 6 — Restore Database from Backup  (Migration Only)")
body(
    "If moving from an existing installation, restore your guru rules, saved filters, and "
    "paper trade history from the GitHub backup."
)
body("1. Find the latest backup file on GitHub:")
code("https://github.com/kishor777/SP500/tree/main/backups")
body("2. Download the most recent backup_YYYY-MM-DD.sql file, then run:")
code("python C:\\SP500\\backup.py --restore C:\\SP500\\backups\\backup_2026-05-29.sql")
body("You will see a count of rows restored for each table.")
note("If this is a fresh installation with no prior data, skip Step 6. The app creates all tables automatically.", "tip")

h2("Step 7 — Start the App")
code("python C:\\SP500\\sp500_viewer.py")
body("You should see:")
code(
    "Loaded 503 tickers from MySQL\n"
    "Loaded history for 503 tickers from MySQL\n"
    " * Running on http://127.0.0.1:5050"
)
body("Open your browser and go to:")
code("http://localhost:5050")
note(
    "On first run, a 2-year history backfill starts in the background. "
    "The app is fully usable immediately — charts show less history until backfill completes (10–30 min).",
    "note"
)

h2("Step 8 — Keep the App Running  (Optional)")
body("By default, closing PowerShell stops the app. To keep it running at startup:")
body("Option A — Run minimised:")
code("Start-Process python -ArgumentList \"C:\\SP500\\sp500_viewer.py\" -WindowStyle Minimized")
body("Option B — Windows Task Scheduler:")
bullet("Open Task Scheduler (search in Start menu)")
bullet("Create Basic Task → name: SP500 App")
bullet("Trigger: When the computer starts")
bullet("Action: Start a program — python  /  Arguments: C:\\SP500\\sp500_viewer.py  /  Start in: C:\\SP500")

h2("3.2  Updating the App (Pulling New Code)")
code(
    "cd C:\\SP500\n"
    "git pull origin main\n"
    "pip install -r requirements.txt"
)
body("Then restart the app. Database tables update automatically on startup.")

h2("3.3  Taking a Manual Backup")
code("python C:\\SP500\\backup.py")
body("This creates backups\\backup_YYYY-MM-DD.sql. Commit and push to GitHub:")
code(
    "git add backups\\\n"
    "git commit -m \"Data backup $(Get-Date -Format yyyy-MM-dd)\"\n"
    "git push"
)

h2("3.4  File Locations Quick Reference")
table(
    ["File", "Location", "In GitHub?"],
    [
        ["App code",         "C:\\SP500\\sp500_viewer.py",          "Yes"],
        ["Configuration",    "C:\\SP500\\config.py",                "Yes"],
        ["Credentials",      "C:\\SP500\\.env",                     "NO — keep private"],
        ["DB connection",    "C:\\SP500\\db_config.py",             "NO — keep private"],
        ["Sentiment module", "C:\\SP500\\sentiment_engine.py",      "Yes"],
        ["Backup script",    "C:\\SP500\\backup.py",                "Yes"],
        ["SQL backups",      "C:\\SP500\\backups\\",                "Yes"],
        ["Requirements",     "C:\\SP500\\requirements.txt",         "Yes"],
    ],
    col_widths=[4, 7, 5.5],
)

h2("3.5  Troubleshooting")
table(
    ["Error / Symptom", "Cause", "Fix"],
    [
        ["ModuleNotFoundError",                "Missing package",          "pip install -r requirements.txt"],
        ["Tables empty, no data",              "Fresh DB, backfill running","Wait 10–30 min, or restore from backup"],
        ["Access denied for user 'root'",      "Wrong password in .env",   "Check DB_PASSWORD matches MySQL root password"],
        ["Charts show no candles",             "No intraday data",         "Run during NYSE market hours (9:30–16:00 ET weekdays)"],
        ["StockTwits empty in logs",           "Corporate proxy blocking", "Normal — app falls back to Yahoo Finance RSS headlines"],
        ["Sentiment shows 'no data yet'",      "First run",                "Click ⚡ Run Sentiment on the Recommendations page"],
        ["Port 5050 already in use",           "Another app on same port", "Run: netstat -ano | findstr :5050 then taskkill /PID <n> /F"],
    ],
    col_widths=[5, 4.5, 7],
)

# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
divider()
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"SP500 App — Maintainer Reference Guide  ·  Generated {date.today().isoformat()}  ·  github.com/kishor777/SP500")
run.font.size  = Pt(8)
run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

# ── Save ──────────────────────────────────────────────────────────────────────
out = f"d:/Apps/SP500_Maintainer_Guide_{date.today().isoformat()}.docx"
doc.save(out)
print(f"Saved: {out}")
